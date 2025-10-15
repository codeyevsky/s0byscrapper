import time
import json
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.service import Service as ChromeService
from selenium.webdriver.chrome.options import Options
from selenium.common.exceptions import TimeoutException, NoSuchElementException
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.colors import black
from datetime import datetime


class TrendyolScraper:
    def __init__(self, headless=False, max_comments=None):
        """
        Trendyol scraper başlatıcı

        Args:
            headless (bool): Tarayıcıyı arka planda çalıştırma
            max_comments (int): Maksimum çekilecek yorum sayısı (None = tümü)
        """
        self.driver = None
        self.headless = headless
        self.max_comments = max_comments
        self.comments = []
        self.reviews = []
        self.product_info = {}

    def setup_driver(self):
        """Selenium WebDriver ayarları - Selenium Manager otomatik driver indirir"""
        chrome_options = Options()

        if self.headless:
            chrome_options.add_argument("--headless=new")

        chrome_options.add_argument("--no-sandbox")
        chrome_options.add_argument("--disable-dev-shm-usage")
        chrome_options.add_argument("--disable-blink-features=AutomationControlled")
        chrome_options.add_argument("--disable-gpu")
        chrome_options.add_argument("--window-size=1920,1080")
        chrome_options.add_argument("user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36")

        # Selenium 4.6+ otomatik driver yönetimi kullanır
        self.driver = webdriver.Chrome(options=chrome_options)
        self.driver.maximize_window()

    def scrape_product(self, url, scrape_mode='comments'):
        """
        Ürün bilgilerini ve yorumları/değerlendirmeleri çeker

        Args:
            url (str): Trendyol ürün URL'si
            scrape_mode (str): 'comments' (yorumlar) veya 'reviews' (değerlendirmeler)

        Returns:
            dict: Ürün bilgileri ve yorumlar/değerlendirmeler
        """
        try:
            self.setup_driver()
            print(f"URL açılıyor: {url}")
            print(f"Scraping modu: {scrape_mode}")
            self.driver.get(url)

            # Sayfanın yüklenmesini bekle
            time.sleep(3)

            # Ürün bilgilerini al
            self._extract_product_info()

            # Yorumlar sekmesine git
            self._navigate_to_comments()

            # Scraping moduna göre veri çek
            if scrape_mode == 'reviews':
                # Değerlendirmeleri çek
                self._extract_reviews_from_html()
                return {
                    'product_info': self.product_info,
                    'reviews': self.reviews,
                    'total_reviews': len(self.reviews),
                    'scrape_mode': 'reviews'
                }
            else:
                # Yorumları çek (varsayılan)
                self._extract_comments()
                return {
                    'product_info': self.product_info,
                    'comments': self.comments,
                    'total_comments': len(self.comments),
                    'scrape_mode': 'comments'
                }

        except Exception as e:
            print(f"Hata oluştu: {str(e)}")
            raise
        finally:
            if self.driver:
                self.driver.quit()

    def _extract_json_ld(self):
        """Sayfadaki JSON-LD verilerini çeker"""
        try:
            script_tags = self.driver.find_elements(By.CSS_SELECTOR, 'script[type="application/ld+json"]')
            json_ld_data = []

            for script in script_tags:
                try:
                    data = json.loads(script.get_attribute('innerHTML'))
                    json_ld_data.append(data)
                except json.JSONDecodeError:
                    continue

            return json_ld_data
        except Exception as e:
            print(f"JSON-LD çekilirken hata: {str(e)}")
            return []

    def _extract_product_info(self):
        """Ürün bilgilerini JSON-LD'den çeker"""
        try:
            # JSON-LD verilerini al
            json_ld_data = self._extract_json_ld()

            # Product verilerini bul
            product_data = None
            for data in json_ld_data:
                if isinstance(data, dict) and data.get('@type') == 'Product':
                    product_data = data
                    break

            if product_data:
                # Ürün adı
                self.product_info['name'] = product_data.get('name', 'Ürün adı bulunamadı')

                # Puan
                rating_data = product_data.get('aggregateRating', {})
                if isinstance(rating_data, dict):
                    rating_value = rating_data.get('ratingValue', 'N/A')
                    rating_count = rating_data.get('ratingCount', 0)
                    self.product_info['rating'] = f"{rating_value} ({rating_count} değerlendirme)"
                else:
                    self.product_info['rating'] = 'Puan bulunamadı'

                print(f"Ürün bilgisi: {self.product_info['name']}")
            else:
                # Fallback: CSS selector kullan
                print("JSON-LD'de ürün bulunamadı, CSS selector kullanılıyor...")
                self._extract_product_info_fallback()

        except Exception as e:
            print(f"Ürün bilgisi çekilirken hata: {str(e)}")
            self._extract_product_info_fallback()

    def _extract_product_info_fallback(self):
        """Ürün bilgilerini CSS selector ile çeker (fallback)"""
        try:
            # Ürün adı - info-title-row class'ından
            try:
                # Önce info-title-row class'ını dene
                title_row = self.driver.find_element(By.CSS_SELECTOR, ".info-title-row, [class*='info-title-row']")

                # Ürün adı (başlık)
                try:
                    product_name = title_row.find_element(By.CSS_SELECTOR, "h1, [class*='title']").text
                    self.product_info['name'] = product_name
                except:
                    self.product_info['name'] = title_row.text.split('\n')[0] if title_row.text else "Ürün adı bulunamadı"

            except:
                # info-title-row bulunamazsa, eski selector'ları dene
                try:
                    product_name = self.driver.find_element(By.CSS_SELECTOR, "h1.pr-new-br, h1").text
                    self.product_info['name'] = product_name
                except:
                    self.product_info['name'] = "Ürün adı bulunamadı"

            # Puan - rate class'ından
            try:
                rating = self.driver.find_element(By.CSS_SELECTOR, ".rate, [class*='rate'], div.rating-score span, span[class*='rating']").text
                self.product_info['rating'] = rating
            except:
                self.product_info['rating'] = "Puan bulunamadı"

            print(f"Ürün bilgisi (fallback): {self.product_info['name']}")

        except Exception as e:
            print(f"Ürün bilgisi (fallback) çekilirken hata: {str(e)}")

    def _navigate_to_comments(self):
        """Yorumlar bölümüne gider"""
        try:
            # Sayfayı aşağı kaydır
            time.sleep(2)
            self.driver.execute_script("window.scrollTo(0, document.body.scrollHeight/2);")
            time.sleep(2)

            # Yorumlar sekmesini bul ve tıkla
            try:
                # Çeşitli selector'ları dene
                possible_selectors = [
                    "//a[contains(@href, '#comments')]",
                    "//div[contains(text(), 'Değerlendirmeler')]",
                    "//button[contains(text(), 'Değerlendirmeler')]",
                    "//a[contains(text(), 'Yorumlar')]"
                ]

                clicked = False
                for selector in possible_selectors:
                    try:
                        comments_tab = WebDriverWait(self.driver, 5).until(
                            EC.element_to_be_clickable((By.XPATH, selector))
                        )
                        self.driver.execute_script("arguments[0].click();", comments_tab)
                        time.sleep(2)
                        print("Yorumlar sekmesine geçildi")
                        clicked = True
                        break
                    except:
                        continue

                if not clicked:
                    print("Yorumlar sekmesi bulunamadı, JSON-LD'den yorumlar çekilecek")

            except Exception as e:
                print(f"Yorumlar sekmesi hatası: {str(e)}")

        except Exception as e:
            print(f"Yorumlar bölümüne geçilirken hata: {str(e)}")

    def _extract_comments(self):
        """Tüm yorumları çeker - önce JSON-LD'den, sonra HTML'den"""
        try:
            print("Yorumlar çekiliyor...")

            # Eğer max_comments belirtilmişse direkt HTML'den çek
            # Çünkü JSON-LD'de sınırlı sayıda yorum var (genellikle ~20)
            if self.max_comments:
                print(f"Maksimum {self.max_comments} yorum istendiği için HTML'den çekilecek...")
                self._extract_comments_from_html()
                return

            # max_comments belirtilmemişse önce JSON-LD'den dene
            json_ld_data = self._extract_json_ld()

            # Review verilerini bul
            reviews_found = False
            for data in json_ld_data:
                if isinstance(data, dict):
                    # Product içindeki review'ları kontrol et
                    if data.get('@type') == 'Product' and 'review' in data:
                        reviews = data.get('review', [])
                        if not isinstance(reviews, list):
                            reviews = [reviews]

                        for review in reviews:
                            if isinstance(review, dict) and review.get('@type') == 'Review':
                                comment_data = {}

                                # Kullanıcı adı
                                author = review.get('author', {})
                                if isinstance(author, dict):
                                    comment_data['user'] = author.get('name', 'Anonim')
                                else:
                                    comment_data['user'] = str(author) if author else 'Anonim'

                                # Yorum metni
                                comment_data['comment'] = review.get('reviewBody', '')

                                # Tarih
                                date_published = review.get('datePublished', '')
                                comment_data['date'] = date_published if date_published else 'Tarih yok'

                                # Tekrar kontrolü - SADECE yorum metnine bak
                                if comment_data['comment']:
                                    is_duplicate = False
                                    for existing_comment in self.comments:
                                        # Sadece yorum metni aynıysa tekrar
                                        if existing_comment.get('comment') == comment_data['comment']:
                                            is_duplicate = True
                                            break

                                    if not is_duplicate:
                                        self.comments.append(comment_data)
                                        reviews_found = True

            if reviews_found:
                print(f"JSON-LD'den {len(self.comments)} yorum çekildi")
            else:
                print("JSON-LD'de yorum bulunamadı, HTML'den çekilecek...")
                self._extract_comments_from_html()

        except Exception as e:
            print(f"Yorumlar çekilirken hata: {str(e)}")
            print("HTML'den yorumlar çekilmeye çalışılıyor...")
            self._extract_comments_from_html()

    def _extract_comments_from_html(self):
        """HTML'den yorumları çeker (fallback)"""
        try:
            # Tüm yorumları yüklemek için "Daha Fazla Göster" butonuna tıkla
            self._load_all_comments()

            print(f"\nYorumlar işlenmeye başlanıyor...")
            print(f"Hedef yorum sayısı: {self.max_comments if self.max_comments else 'Tümü'}")

            # Sadece class="review" olan div'leri seç
            possible_selectors = [
                "div.review"
            ]

            comment_elements = []
            for selector in possible_selectors:
                comment_elements = self.driver.find_elements(By.CSS_SELECTOR, selector)
                if comment_elements:
                    print(f"'{selector}' ile {len(comment_elements)} yorum elementi bulundu")
                    break

            if not comment_elements:
                print("HTML'de yorum bulunamadı")
                return

            # İlk 2 elementi atla (genellikle filtre/sıralama elementleri)
            comment_elements = comment_elements[2:] if len(comment_elements) > 2 else comment_elements
            print(f"İşlenecek yorum elementi sayısı (filtrelemeden sonra): {len(comment_elements)}")

            for idx, comment_elem in enumerate(comment_elements, 1):
                # Maksimum yorum sayısına ulaşıldıysa dur
                if self.max_comments and len(self.comments) >= self.max_comments:
                    print(f"\n✓ Hedef yorum sayısına ulaşıldı: {len(self.comments)}/{self.max_comments}")
                    break

                try:
                    # Yorumu görünür hale getir (scroll)
                    self.driver.execute_script("arguments[0].scrollIntoView({behavior: 'auto', block: 'center'});", comment_elem)
                    time.sleep(0.2)

                    comment_data = {}

                    # "Devamını oku" butonunu kontrol et ve tıkla
                    try:
                        # Önce XPath ile text içeriğine göre ara
                        try:
                            read_more_button = comment_elem.find_element(By.XPATH, ".//a[contains(text(), 'Devamını oku')] | .//button[contains(text(), 'Devamını oku')] | .//span[contains(text(), 'Devamını oku')]")
                        except:
                            # Yoksa class'a göre ara
                            read_more_button = comment_elem.find_element(By.CSS_SELECTOR, "a[class*='read-more'], button[class*='read-more'], [class*='show-more'], [class*='devamini-oku']")

                        if read_more_button and read_more_button.is_displayed():
                            self.driver.execute_script("arguments[0].click();", read_more_button)
                            time.sleep(0.5)
                            print(f"Yorum {idx} için 'Devamını oku' butonuna tıklandı")
                    except:
                        pass  # Buton yoksa devam et

                    # Kullanıcı adı - class="name"
                    try:
                        user_name = comment_elem.find_element(By.CSS_SELECTOR, ".name").text
                        # Tüm whitespace'leri tek boşluğa çevir
                        comment_data['user'] = ' '.join(user_name.split()) if user_name else "Anonim"
                    except:
                        comment_data['user'] = "Anonim"

                    # Yorum metni - SADECE review-comment class'ını kullan
                    try:
                        review_comment = comment_elem.find_element(By.CSS_SELECTOR, "span.review-comment")
                        comment_data['comment'] = review_comment.text.strip()
                    except:
                        comment_data['comment'] = ""

                    # Tarih - class="date"
                    try:
                        date = comment_elem.find_element(By.CSS_SELECTOR, ".date").text
                        # Tüm whitespace'leri (newline, tab, vs.) tek boşluğa çevir
                        comment_data['date'] = ' '.join(date.split()) if date else "Tarih yok"
                    except:
                        comment_data['date'] = "Tarih yok"

                    # Tekrar kontrolü - SADECE yorum metnine bak (kullanıcı adına BAKMA!)
                    if comment_data['comment']:
                        is_duplicate = False
                        for existing_comment in self.comments:
                            # Sadece yorum metni aynıysa tekrar, kullanıcı farklı olabilir
                            if existing_comment.get('comment') == comment_data['comment']:
                                is_duplicate = True
                                print(f"Yorum {idx} tekrar ediyor (aynı metin), atlanıyor")
                                break

                        if not is_duplicate:
                            self.comments.append(comment_data)
                            print(f"✓ Yorum {idx} eklendi (Toplam: {len(self.comments)})")
                    else:
                        print(f"Yorum {idx} boş veya çekilemedi, atlanıyor")

                except Exception as e:
                    print(f"✗ Yorum {idx} çekilirken hata: {str(e)}")
                    continue

            if self.max_comments and len(self.comments) < self.max_comments:
                print(f"\n⚠ Uyarı: Hedef yorum sayısına ulaşılamadı. İstenen: {self.max_comments}, Çekilen: {len(self.comments)}")
                print(f"Toplam {len(comment_elements)} element işlendi, {len(self.comments)} benzersiz yorum bulundu")
            else:
                print(f"\n✓ HTML'den toplam {len(self.comments)} yorum başarıyla çekildi")

        except Exception as e:
            print(f"HTML'den yorumlar çekilirken hata: {str(e)}")

    def _load_all_comments(self):
        """INFINITE SCROLL: Sayfayı scroll ederek tüm yorumları yükler"""
        max_scrolls = 200  # Maksimum scroll sayısı
        scrolls = 0
        no_new_comments_count = 0  # Yeni yorum gelmeme sayacı

        print("Infinite scroll ile yorumlar yükleniyor...")

        while scrolls < max_scrolls:
            try:
                # Mevcut yorum sayısını al
                previous_count = len(self.driver.find_elements(By.CSS_SELECTOR, "div.review"))

                # Sayfanın en altına scroll yap
                self.driver.execute_script("window.scrollTo(0, document.body.scrollHeight);")
                time.sleep(1.5)  # Yorumların yüklenmesi için bekle

                # Yeni yorum sayısını al
                current_count = len(self.driver.find_elements(By.CSS_SELECTOR, "div.review"))

                scrolls += 1

                # Yeni yorum geldi mi?
                if current_count > previous_count:
                    new_comments = current_count - previous_count
                    print(f"Scroll #{scrolls}: {new_comments} yeni yorum yüklendi (Toplam: {current_count})")
                    no_new_comments_count = 0  # Sayacı sıfırla
                else:
                    no_new_comments_count += 1
                    print(f"Scroll #{scrolls}: Yeni yorum yok (Toplam: {current_count})")

                # 5 kez üst üste yeni yorum gelmediyse dur
                if no_new_comments_count >= 5:
                    print(f"\n5 kez üst üste yeni yorum gelmedi. Tüm yorumlar yüklendi (Toplam: {current_count})")
                    break

                # Hedef sayıya ulaşıldıysa dur (buffer ile)
                if self.max_comments and current_count >= self.max_comments * 2:
                    print(f"\nYeterli yorum yüklendi ({current_count}). Hedef: {self.max_comments}")
                    break

            except Exception as e:
                print(f"Scroll hatası: {str(e)}")
                break

        final_count = len(self.driver.find_elements(By.CSS_SELECTOR, "div.review"))
        print(f"\nScroll tamamlandı. Toplam {final_count} yorum yüklendi ({scrolls} scroll)")

    def _load_all_reviews(self):
        """INFINITE SCROLL: review-list-scroll-container içinde scroll ederek tüm değerlendirmeleri yükler"""
        max_scrolls = 200  # Maksimum scroll sayısı
        scrolls = 0
        no_new_reviews_count = 0  # Yeni değerlendirme gelmeme sayacı

        print("Infinite scroll ile değerlendirmeler yükleniyor...")

        try:
            # review-list-scroll-container'ı bul
            scroll_container = self.driver.find_element(By.CSS_SELECTOR, ".review-list-scroll-container")
        except:
            print("review-list-scroll-container bulunamadı, normal scroll kullanılacak")
            scroll_container = None

        while scrolls < max_scrolls:
            try:
                # Mevcut review sayısını al
                previous_count = len(self.driver.find_elements(By.CSS_SELECTOR, ".review-list .review"))

                # Scroll yap
                if scroll_container:
                    # Container içinde scroll
                    self.driver.execute_script("arguments[0].scrollTop = arguments[0].scrollHeight", scroll_container)
                else:
                    # Normal sayfa scroll
                    self.driver.execute_script("window.scrollTo(0, document.body.scrollHeight);")

                time.sleep(1.5)  # Değerlendirmelerin yüklenmesi için bekle

                # Yeni review sayısını al
                current_count = len(self.driver.find_elements(By.CSS_SELECTOR, ".review-list .review"))

                scrolls += 1

                # Yeni review geldi mi?
                if current_count > previous_count:
                    new_reviews = current_count - previous_count
                    print(f"Scroll #{scrolls}: {new_reviews} yeni değerlendirme yüklendi (Toplam: {current_count})")
                    no_new_reviews_count = 0  # Sayacı sıfırla
                else:
                    no_new_reviews_count += 1
                    print(f"Scroll #{scrolls}: Yeni değerlendirme yok (Toplam: {current_count})")

                # 5 kez üst üste yeni review gelmediyse dur
                if no_new_reviews_count >= 5:
                    print(f"\n5 kez üst üste yeni değerlendirme gelmedi. Tümü yüklendi (Toplam: {current_count})")
                    break

                # Hedef sayıya ulaşıldıysa dur (buffer ile)
                if self.max_comments and current_count >= self.max_comments * 2:
                    print(f"\nYeterli değerlendirme yüklendi ({current_count}). Hedef: {self.max_comments}")
                    break

            except Exception as e:
                print(f"Scroll hatası: {str(e)}")
                break

        final_count = len(self.driver.find_elements(By.CSS_SELECTOR, ".review-list .review"))
        print(f"\nScroll tamamlandı. Toplam {final_count} değerlendirme yüklendi ({scrolls} scroll)")

    def _extract_reviews_from_html(self):
        """review-container yapısından değerlendirmeleri çeker"""
        try:
            # Tüm değerlendirmeleri yüklemek için scroll yap
            self._load_all_reviews()

            print(f"\nDeğerlendirmeler işlenmeye başlanıyor...")
            print(f"Hedef değerlendirme sayısı: {self.max_comments if self.max_comments else 'Tümü'}")

            # .review-list .review elementlerini bul
            review_elements = self.driver.find_elements(By.CSS_SELECTOR, ".review-list .review")

            if not review_elements:
                print("Değerlendirme bulunamadı")
                return

            print(f"İşlenecek değerlendirme elementi sayısı: {len(review_elements)}")

            for idx, review_elem in enumerate(review_elements, 1):
                # Maksimum review sayısına ulaşıldıysa dur
                if self.max_comments and len(self.reviews) >= self.max_comments:
                    print(f"\n✓ Hedef değerlendirme sayısına ulaşıldı: {len(self.reviews)}/{self.max_comments}")
                    break

                try:
                    # Review'ı görünür hale getir (scroll)
                    self.driver.execute_script("arguments[0].scrollIntoView({behavior: 'auto', block: 'center'});", review_elem)
                    time.sleep(0.2)

                    review_data = {}

                    # "Devamını oku" butonunu kontrol et ve tıkla
                    try:
                        read_more_button = review_elem.find_element(By.XPATH, ".//a[contains(text(), 'Devamını oku')] | .//button[contains(text(), 'Devamını oku')] | .//span[contains(text(), 'Devamını oku')]")
                        if read_more_button and read_more_button.is_displayed():
                            self.driver.execute_script("arguments[0].click();", read_more_button)
                            time.sleep(0.5)
                            print(f"Değerlendirme {idx} için 'Devamını oku' butonuna tıklandı")
                    except:
                        pass  # Buton yoksa devam et

                    # item-header > seller
                    try:
                        seller = review_elem.find_element(By.CSS_SELECTOR, ".item-header .seller")
                        review_data['seller'] = ' '.join(seller.text.split()).strip() if seller.text else "Satıcı bulunamadı"
                    except:
                        review_data['seller'] = "Satıcı bulunamadı"

                    # item-header > product
                    try:
                        product = review_elem.find_element(By.CSS_SELECTOR, ".item-header .product")
                        review_data['product'] = ' '.join(product.text.split()).strip() if product.text else "Ürün bulunamadı"
                    except:
                        review_data['product'] = "Ürün bulunamadı"

                    # review-info > name-wrapper > comment
                    try:
                        comment = review_elem.find_element(By.CSS_SELECTOR, ".review-info .name-wrapper .comment")
                        review_data['comment'] = comment.text.strip() if comment.text else ""
                    except:
                        review_data['comment'] = ""

                    # review-info > review-info-detail > name
                    try:
                        name = review_elem.find_element(By.CSS_SELECTOR, ".review-info .review-info-detail .name")
                        review_data['name'] = ' '.join(name.text.split()).strip() if name.text else "Anonim"
                    except:
                        review_data['name'] = "Anonim"

                    # review-info > review-info-detail > date
                    try:
                        date = review_elem.find_element(By.CSS_SELECTOR, ".review-info .review-info-detail .date")
                        review_data['date'] = ' '.join(date.text.split()).strip() if date.text else "Tarih yok"
                    except:
                        review_data['date'] = "Tarih yok"

                    # Tekrar kontrolü - SADECE comment text'e bak
                    if review_data['comment']:
                        is_duplicate = False
                        for existing_review in self.reviews:
                            if existing_review.get('comment') == review_data['comment']:
                                is_duplicate = True
                                print(f"Değerlendirme {idx} tekrar ediyor (aynı metin), atlanıyor")
                                break

                        if not is_duplicate:
                            self.reviews.append(review_data)
                            print(f"✓ Değerlendirme {idx} eklendi (Toplam: {len(self.reviews)})")
                    else:
                        print(f"Değerlendirme {idx} boş veya çekilemedi, atlanıyor")

                except Exception as e:
                    print(f"✗ Değerlendirme {idx} çekilirken hata: {str(e)}")
                    continue

            if self.max_comments and len(self.reviews) < self.max_comments:
                print(f"\n⚠ Uyarı: Hedef değerlendirme sayısına ulaşılamadı. İstenen: {self.max_comments}, Çekilen: {len(self.reviews)}")
                print(f"Toplam {len(review_elements)} element işlendi, {len(self.reviews)} benzersiz değerlendirme bulundu")
            else:
                print(f"\n✓ Toplam {len(self.reviews)} değerlendirme başarıyla çekildi")

        except Exception as e:
            print(f"Değerlendirmeler çekilirken hata: {str(e)}")

    def _extract_rating(self, rating_class):
        """Yıldız puanını class'tan çıkarır"""
        try:
            if 'full' in rating_class.lower():
                return "5"
            elif '4' in rating_class:
                return "4"
            elif '3' in rating_class:
                return "3"
            elif '2' in rating_class:
                return "2"
            elif '1' in rating_class:
                return "1"
            return "N/A"
        except:
            return "N/A"

    def export_to_word(self, filename="trendyol_yorumlar.docx"):
        """
        Yorumları veya değerlendirmeleri Word dosyasına aktarır

        Args:
            filename (str): Çıktı dosya adı
        """
        # Reviews var mı kontrol et
        has_reviews = len(self.reviews) > 0
        has_comments = len(self.comments) > 0

        if not has_reviews and not has_comments:
            print("Henüz veri çekilmedi!")
            return

        doc = Document()

        # Başlık
        if has_reviews:
            title = doc.add_heading('Trendyol Ürün Değerlendirmeleri', 0)
        else:
            title = doc.add_heading('Trendyol Ürün Yorumları', 0)
        title.alignment = 1  # Ortalanmış

        # Ürün bilgileri
        doc.add_heading('Ürün Bilgileri', level=1)
        for key, value in self.product_info.items():
            p = doc.add_paragraph()
            p.add_run(f"{key.capitalize()}: ").bold = True
            p.add_run(str(value))

        doc.add_paragraph()

        # Reviews varsa onu export et
        if has_reviews:
            doc.add_heading(f'Toplam Değerlendirme Sayısı: {len(self.reviews)}', level=2)
            doc.add_paragraph()

            # Değerlendirmeler
            doc.add_heading('Değerlendirmeler', level=1)

            for idx, review in enumerate(self.reviews, 1):
                # Değerlendirme başlığı
                doc.add_heading(f'Değerlendirme #{idx}', level=2)

                # Satıcı
                p = doc.add_paragraph()
                p.add_run('Satıcı: ').bold = True
                p.add_run(review.get('seller', 'Bilinmiyor'))

                # Ürün
                p = doc.add_paragraph()
                p.add_run('Ürün: ').bold = True
                p.add_run(review.get('product', 'Bilinmiyor'))

                # Kullanıcı bilgisi
                p = doc.add_paragraph()
                p.add_run('Kullanıcı: ').bold = True
                p.add_run(review.get('name', 'Anonim'))

                # Tarih
                p = doc.add_paragraph()
                p.add_run('Tarih: ').bold = True
                p.add_run(review.get('date', 'Bilinmiyor'))

                # Yorum metni
                p = doc.add_paragraph()
                p.add_run('Değerlendirme: ').bold = True
                doc.add_paragraph(review.get('comment', ''))

                # Ayırıcı
                doc.add_paragraph('_' * 80)

        # Comments varsa onu export et
        else:
            doc.add_heading(f'Toplam Yorum Sayısı: {len(self.comments)}', level=2)
            doc.add_paragraph()

            # Yorumlar
            doc.add_heading('Yorumlar', level=1)

            for idx, comment in enumerate(self.comments, 1):
                # Yorum başlığı
                doc.add_heading(f'Yorum #{idx}', level=2)

                # Kullanıcı bilgisi
                p = doc.add_paragraph()
                p.add_run('Kullanıcı: ').bold = True
                p.add_run(comment.get('user', 'Anonim'))

                # Tarih
                p = doc.add_paragraph()
                p.add_run('Tarih: ').bold = True
                p.add_run(comment.get('date', 'Bilinmiyor'))

                # Yorum metni
                p = doc.add_paragraph()
                p.add_run('Yorum: ').bold = True
                doc.add_paragraph(comment.get('comment', ''))

                # Ayırıcı
                doc.add_paragraph('_' * 80)

        doc.save(filename)
        print(f"Word dosyası oluşturuldu: {filename}")

    def export_to_pdf(self, filename="trendyol_yorumlar.pdf"):
        """
        Yorumları PDF dosyasına aktarır

        Args:
            filename (str): Çıktı dosya adı
        """
        if not self.comments:
            print("Henüz yorum çekilmedi!")
            return

        # PDF oluştur
        doc = SimpleDocTemplate(filename, pagesize=A4)
        story = []
        styles = getSampleStyleSheet()

        # Özel stiller
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=black,
            spaceAfter=30,
            alignment=TA_LEFT
        )

        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=16,
            textColor=black,
            spaceAfter=12,
            spaceBefore=12
        )

        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=11,
            leading=14,
            alignment=TA_LEFT
        )

        # Başlık
        story.append(Paragraph("Trendyol Ürün Yorumları", title_style))
        story.append(Spacer(1, 0.2*inch))

        # Ürün bilgileri
        story.append(Paragraph("Ürün Bilgileri", heading_style))
        for key, value in self.product_info.items():
            text = f"<b>{key.capitalize()}:</b> {value}"
            story.append(Paragraph(text, normal_style))

        story.append(Spacer(1, 0.2*inch))
        story.append(Paragraph(f"<b>Toplam Yorum Sayısı:</b> {len(self.comments)}", normal_style))
        story.append(Spacer(1, 0.3*inch))

        # Yorumlar
        story.append(Paragraph("Yorumlar", heading_style))
        story.append(Spacer(1, 0.2*inch))

        for idx, comment in enumerate(self.comments, 1):
            # Yorum başlığı
            story.append(Paragraph(f"<b>Yorum #{idx}</b>", heading_style))

            # Kullanıcı
            text = f"<b>Kullanıcı:</b> {comment.get('user', 'Anonim')}"
            story.append(Paragraph(text, normal_style))

            # Tarih
            text = f"<b>Tarih:</b> {comment.get('date', 'Bilinmiyor')}"
            story.append(Paragraph(text, normal_style))

            # Yorum metni
            comment_text = comment.get('comment', '').replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            text = f"<b>Yorum:</b> {comment_text}"
            story.append(Paragraph(text, normal_style))

            story.append(Spacer(1, 0.2*inch))
            story.append(Paragraph("_" * 100, normal_style))
            story.append(Spacer(1, 0.2*inch))

        # PDF'i oluştur
        doc.build(story)
        print(f"PDF dosyası oluşturuldu: {filename}")


def main():
    """Ana fonksiyon - örnek kullanım"""
    # Ürün mü Mağaza mı seçimi (İLK SORU)
    print("Neyin değerlendirmelerini çekmek istersiniz?")
    print("1) Ürün Yorumları")
    print("2) Mağaza Değerlendirmeleri")
    mode_choice = input("Seçiminiz (1 veya 2): ").strip()

    scrape_mode = 'reviews' if mode_choice == '2' else 'comments'

    # Trendyol ürün URL'si
    url = input("\nTrendyol ürün URL'sini girin: ")

    # Maksimum yorum/değerlendirme sayısı
    data_type = "değerlendirme" if scrape_mode == 'reviews' else "yorum"
    max_comments_input = input(f"Maksimum kaç {data_type} çekmek istersiniz? (Tümü için Enter'a basın): ").strip()
    max_comments = int(max_comments_input) if max_comments_input else None

    # Scraper başlat
    scraper = TrendyolScraper(headless=False, max_comments=max_comments)  # headless=True yaparak arka planda çalıştırabilirsiniz

    try:
        # Verileri çek
        result = scraper.scrape_product(url, scrape_mode=scrape_mode)

        print(f"\n{'='*50}")
        print(f"Ürün: {result['product_info'].get('name', 'Bilinmiyor')}")

        if scrape_mode == 'reviews':
            print(f"Toplam {result['total_reviews']} mağaza değerlendirmesi çekildi")
        else:
            print(f"Toplam {result['total_comments']} ürün yorumu çekildi")

        print(f"{'='*50}\n")

        # Dosyaya kaydet (sadece Word formatı)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"trendyol_{scrape_mode}_{timestamp}.docx"
        scraper.export_to_word(filename)

        print("\nİşlem tamamlandı!")

    except Exception as e:
        print(f"Hata: {str(e)}")


if __name__ == "__main__":
    main()
